from __future__ import annotations

import json
import math
import os
import re
import shutil
import tempfile
import zipfile
from collections import defaultdict
from dataclasses import dataclass
from datetime import date, datetime, timedelta
from pathlib import Path
from typing import Any
from urllib import request
from urllib.error import HTTPError, URLError

from .bot_taxonomy import load_bot_taxonomy
from .index_filtering import filter_index_options
from .support import DashboardQueryError, local_day_to_utc_bounds


ROOT_DIR = Path(__file__).resolve().parent.parent
REPORT_TEMPLATE_PATH = ROOT_DIR / "2026-04-22-qbedding-traffic-report-client.docx"
REPORT_OUTPUT_DIR = ROOT_DIR / "output" / "current" / "reports"
REPORT_MANIFEST_PATH = REPORT_OUTPUT_DIR / "_manifest.json"
REPORT_REMOTE_PATH = "/internal/search/es"
REPORT_TIMEZONE = "Asia/Shanghai"
REPORT_LOCAL_CONFIG_PATH = ROOT_DIR / "reporting.local.json"
REPORT_RETENTION_PER_CUSTOMER = 5
REPORT_RETENTION_DAYS = 7


@dataclass
class WeekStats:
    label: str
    start: str
    end: str
    days: int
    ai_index: int
    ai_search: int
    ai_training: int
    seo_bot: int

    @property
    def total(self) -> int:
        return self.ai_index + self.ai_search + self.ai_training + self.seo_bot

    @property
    def ai_total(self) -> int:
        return self.ai_index + self.ai_search + self.ai_training

    def daily_avg(self, value: int) -> float:
        return value / self.days if self.days else 0.0


def generate_word_report(
    service,
    customer_name: str | None,
    host: str | None,
    date_from: str,
    date_to: str,
    summary_override: str | None = None,
    llm_sections_override: dict[str, Any] | None = None,
) -> dict[str, str]:
    try:
        from docx import Document
    except ImportError as exc:
        raise DashboardQueryError(
            message="report generation dependency missing: python-docx",
            error_type="report_dependency_missing",
            source="reporting",
            status_code=500,
            extra={"module": "python-docx"},
        ) from exc

    _ensure_template_exists()
    report_data = _collect_report_data(
        service,
        customer_name,
        host,
        date_from,
        date_to,
        include_llm_summary=summary_override is None,
    )
    if summary_override:
        report_data["llm_summary"] = summary_override
    if llm_sections_override:
        report_data["llm_sections"] = llm_sections_override
    REPORT_OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    output_name = _report_filename(report_data["file_label"], date_from, date_to)
    output_path = REPORT_OUTPUT_DIR / output_name
    shutil.copyfile(REPORT_TEMPLATE_PATH, output_path)

    doc = Document(str(output_path))
    _fill_docx_template(doc, report_data)
    doc.save(str(output_path))

    with tempfile.TemporaryDirectory(prefix="report_charts_") as tmp_dir:
        chart_paths = _build_chart_images(report_data, Path(tmp_dir))
        _replace_docx_images(output_path, chart_paths)

    _cleanup_report_outputs(output_path.name, report_data["customer_key"])

    return {
        "filename": output_name,
        "path": str(output_path),
    }


def build_report_summary_context(service, customer_name: str | None, host: str | None, date_from: str, date_to: str) -> dict[str, Any]:
    data = _collect_report_data(
        service,
        customer_name,
        host,
        date_from,
        date_to,
        include_llm_summary=False,
    )
    return {
        "llm_context": data["llm_context"],
        "fallback_summary": data["llm_summary"],
    }


def generate_weekly_comparison(service, customer_name: str | None, host: str | None, date_from: str, date_to: str) -> dict[str, Any]:
    data = _collect_report_data(service, customer_name, host, date_from, date_to)
    return {
        "title_label": data["title_label"],
        "weeks": [
            {
                "label": week.label,
                "start": week.start,
                "end": week.end,
                "days": week.days,
                "total": week.total,
                "daily_avg": round(week.daily_avg(week.total), 1),
                "ai_daily_avg": round(week.daily_avg(week.ai_total), 1),
                "seo_daily_avg": round(week.daily_avg(week.seo_bot), 1),
            }
            for week in data["week_stats"]
        ],
        "note": f"按当前筛选对象在所选时间范围内的全部自然周统计，当前共展示 {len(data['week_stats'])} 个周区间。",
    }


def resolve_report_download_path(file_name: str) -> Path:
    safe_name = Path(file_name).name
    path = (REPORT_OUTPUT_DIR / safe_name).resolve()
    base = REPORT_OUTPUT_DIR.resolve()
    if base not in path.parents and path != base:
        raise DashboardQueryError(
            message="invalid report path",
            error_type="report_download_invalid_path",
            source="reporting",
            status_code=400,
        )
    if not path.exists() or not path.is_file():
        raise DashboardQueryError(
            message="report file not found",
            error_type="report_download_missing",
            source="reporting",
            status_code=404,
        )
    return path


def _ensure_template_exists() -> None:
    if REPORT_TEMPLATE_PATH.exists():
        return
    raise DashboardQueryError(
        message="report template file is missing",
        error_type="report_template_missing",
        source="reporting",
        status_code=500,
        extra={"path": str(REPORT_TEMPLATE_PATH)},
    )


def _report_filename(file_label: str, date_from: str, date_to: str) -> str:
    label = re.sub(r"[^A-Za-z0-9._-]+", "-", (file_label or "report")).strip("-").lower() or "report"
    generated_on = datetime.now().date().isoformat()
    return f"{generated_on}-{label}-traffic-report-{date_from}_to_{date_to}.docx"


def _normalize_customer_key(value: str | None) -> str:
    normalized = re.sub(r"[^A-Za-z0-9._-]+", "-", str(value or "").strip()).strip("-").lower()
    return normalized or "all-customers"


def _load_report_manifest() -> dict[str, dict[str, Any]]:
    if not REPORT_MANIFEST_PATH.exists():
        return {}
    try:
        payload = json.loads(REPORT_MANIFEST_PATH.read_text(encoding="utf-8"))
    except Exception:
        return {}
    if not isinstance(payload, dict):
        return {}
    manifest: dict[str, dict[str, Any]] = {}
    for name, meta in payload.items():
        if not isinstance(name, str) or not isinstance(meta, dict):
            continue
        manifest[name] = {
            "customer_key": _normalize_customer_key(meta.get("customer_key")),
            "generated_at": str(meta.get("generated_at") or "").strip(),
        }
    return manifest


def _save_report_manifest(manifest: dict[str, dict[str, Any]]) -> None:
    REPORT_OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    REPORT_MANIFEST_PATH.write_text(
        json.dumps(manifest, ensure_ascii=False, indent=2, sort_keys=True),
        encoding="utf-8",
    )


def _infer_customer_key_from_filename(name: str) -> str:
    base = Path(name).name
    match = re.match(r"^\d{4}-\d{2}-\d{2}-(.+?)-traffic-report-\d{4}-\d{2}-\d{2}_to_\d{4}-\d{2}-\d{2}\.docx$", base)
    if not match:
        return "all-customers"
    return _normalize_customer_key(match.group(1))


def _delete_report_file(path: Path) -> None:
    try:
        if path.exists() and path.is_file():
            path.unlink()
    except OSError:
        pass


def _cleanup_report_outputs(current_filename: str, customer_key: str) -> None:
    REPORT_OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    manifest = _load_report_manifest()
    current_key = _normalize_customer_key(customer_key)
    manifest[current_filename] = {
        "customer_key": current_key,
        "generated_at": datetime.now().isoformat(timespec="seconds"),
    }

    docx_files = {path.name: path for path in REPORT_OUTPUT_DIR.glob("*.docx") if path.is_file()}
    manifest = {name: meta for name, meta in manifest.items() if name in docx_files}
    now = datetime.now()
    cutoff = now - timedelta(days=REPORT_RETENTION_DAYS)

    for name, path in list(docx_files.items()):
        try:
            modified_at = datetime.fromtimestamp(path.stat().st_mtime)
        except OSError:
            continue
        if modified_at < cutoff:
            _delete_report_file(path)
            docx_files.pop(name, None)
            manifest.pop(name, None)

    grouped: dict[str, list[tuple[datetime, str, Path]]] = defaultdict(list)
    for name, path in docx_files.items():
        meta = manifest.get(name) or {}
        key = _normalize_customer_key(meta.get("customer_key")) if meta else _infer_customer_key_from_filename(name)
        manifest.setdefault(name, {"customer_key": key, "generated_at": ""})
        try:
            modified_at = datetime.fromtimestamp(path.stat().st_mtime)
        except OSError:
            modified_at = datetime.min
        grouped[key].append((modified_at, name, path))

    for key, items in grouped.items():
        items.sort(key=lambda item: (item[0], item[1]), reverse=True)
        for _, name, path in items[REPORT_RETENTION_PER_CUSTOMER:]:
            _delete_report_file(path)
            manifest.pop(name, None)

    _save_report_manifest(manifest)


def _collect_report_data(
    service,
    customer_name: str | None,
    host: str | None,
    date_from: str,
    date_to: str,
    include_llm_summary: bool = True,
) -> dict[str, Any]:
    remote = service.remote_source
    start_utc, end_utc = local_day_to_utc_bounds(date_from, date_to)
    period_indices = filter_index_options(
        remote.list_index_options(start_utc=start_utc, end_utc=end_utc),
        customer_name,
        None,
        date_from=date_from,
        date_to=date_to,
    )
    period_index_names = [str(item.get("value") or "") for item in period_indices if item.get("value")]
    if customer_name and customer_name != "ALL" and not period_index_names:
        raise DashboardQueryError(
            message="no data found for selected report scope",
            error_type="report_scope_empty",
            source="reporting",
            status_code=404,
            extra={"customer_name": customer_name, "host": host, "date_from": date_from, "date_to": date_to},
        )

    host_filters = None
    if host and host != "ALL":
        host_options = remote.list_host_options(
            index_names=period_index_names,
            start_utc=start_utc,
            end_utc=end_utc,
        )
        actual_hosts = [str(item.get("value") or "") for item in host_options if str(item.get("value") or "").strip() and str(item.get("value") or "") != "ALL"]
        if len(actual_hosts) > 1:
            host_filters = [host]
    is_shopify = remote._is_shopify_scope(period_index_names)

    current_label = host if host and host != "ALL" else (customer_name if customer_name and customer_name != "ALL" else "all-customers")
    title_label = host if host and host != "ALL" else (customer_name if customer_name and customer_name != "ALL" else "全部客户")
    customer_key = customer_name if customer_name and customer_name != "ALL" else "all-customers"
    month_label = f"{date.fromisoformat(date_to).year} 年 {date.fromisoformat(date_to).month} 月"

    summary = _query_report_summary(remote, period_index_names, host_filters, start_utc, end_utc)
    rankings = _query_category_rankings(remote, period_index_names, host_filters, start_utc, end_utc)
    top_pages = _query_top_pages(remote, period_index_names, host_filters, start_utc, end_utc, limit=25)
    stage = _query_stage_assessment_v2(remote, period_index_names, host_filters, start_utc, end_utc)
    weekly_days = _query_report_daily(remote, period_index_names, host_filters, start_utc, end_utc)
    week_stats = _compute_weeks_in_range(weekly_days)
    dashboard_window = remote.get_live_dashboard_window(
        index_names=period_index_names,
        host_filters=host_filters,
        start_utc=start_utc,
        end_utc=end_utc,
        top_bots=5,
        top_pages=5,
        include_rankings=False,
    )
    dashboard_view = {
        "cards": {
            "total_requests": int(sum(dashboard_window.get("focused", {}).get(key, 0) for key in ("user_traditional", "user_ai", "user_platform", "user_direct", "ai_search", "ai_training", "ai_index"))),
            "user_requests": int(sum(dashboard_window.get("focused", {}).get(key, 0) for key in ("user_traditional", "user_ai", "user_platform", "user_direct"))),
            "ai_requests": int(sum(dashboard_window.get("focused", {}).get(key, 0) for key in ("ai_search", "ai_training", "ai_index"))),
        }
    }
    dashboard_days = []
    for row in dashboard_window.get("days", []):
        user_total = int(sum(row.get(key, 0) for key in ("user_traditional", "user_ai", "user_platform", "user_direct")))
        ai_total = int(sum(row.get(key, 0) for key in ("ai_search", "ai_training", "ai_index")))
        seo_total = int(row.get("seo_bot", 0))
        dashboard_days.append(
            {
                "date": row.get("date"),
                "user_total": user_total,
                "ai_total": ai_total,
                "seo_total": seo_total,
                "total": user_total + ai_total + seo_total,
                "ai_index": int(row.get("ai_index", 0)),
                "ai_search": int(row.get("ai_search", 0)),
                "ai_training": int(row.get("ai_training", 0)),
            }
        )
    taxonomy_rows = [(entry.bot_name, entry.category) for entry in load_bot_taxonomy().entries]
    llm_context = _summary_prompt_payload(
        title_label=title_label,
        date_from=date_from,
        date_to=date_to,
        summary=summary,
        stage=stage,
    )
    llm_context.update(
        {
            "is_shopify": is_shopify,
            "period_days": _inclusive_days(date_from, date_to),
            "ai_category_rankings": rankings,
            "overall_top_ai_platform": _overall_top_ai_platform(rankings),
            "top_pages": [
                {
                    "page": item["page"],
                    "total": item["total"],
                    "ai_index": item["ai_index"],
                    "ai_search": item["ai_search"],
                    "ai_training": item["ai_training"],
                    "seo_bot": item["seo_bot"],
                }
                for item in top_pages[:5]
            ],
            "daily_breakdown": [
                {
                    "date": row.get("date"),
                    "total": int(row.get("ai_index", 0) + row.get("ai_search", 0) + row.get("ai_training", 0) + row.get("seo_bot", 0)),
                    "ai_index": int(row.get("ai_index", 0)),
                    "ai_search": int(row.get("ai_search", 0)),
                    "ai_training": int(row.get("ai_training", 0)),
                    "seo_bot": int(row.get("seo_bot", 0)),
                }
                for row in summary.get("days", [])
            ],
            "weeks": [
                {
                    "label": week.label,
                    "start": week.start,
                    "end": week.end,
                    "days": week.days,
                    "total": week.total,
                    "ai_index": week.ai_index,
                    "ai_search": week.ai_search,
                    "ai_training": week.ai_training,
                    "seo_bot": week.seo_bot,
                }
                for week in week_stats
            ],
            "dashboard_cards": dashboard_view["cards"],
            "anomalies": _detect_anomalies(
                {
                    "summary": summary,
                    "stage": stage,
                    "dashboard_view": dashboard_view,
                    "week_stats": week_stats,
                    "top_pages": top_pages,
                    "title_label": title_label,
                },
                shopify=is_shopify,
            ),
            "compare_series": dashboard_days,
        }
    )
    llm_summary = (
        generate_llm_summary_from_context(llm_context)
        if include_llm_summary
        else _stage_summary_text(title_label, stage)
    )

    return {
        "title_label": title_label,
        "scope_label": title_label,
        "file_label": current_label,
        "customer_key": customer_key,
        "is_shopify": is_shopify,
        "month_label": month_label,
        "date_from": date_from,
        "date_to": date_to,
        "period_days": _inclusive_days(date_from, date_to),
        "summary": summary,
        "rankings": rankings,
        "top_pages": top_pages,
        "stage": stage,
        "week_stats": week_stats,
        "dashboard_view": dashboard_view,
        "dashboard_days": dashboard_days,
        "taxonomy_rows": taxonomy_rows,
        "selected_week_note": _period_week_note(date_from, date_to),
        "final_note": "本报告由 Deeplumen 提供 | 数据来源：Agentic Page访问日志" if is_shopify else "本报告由 Deeplumen 提供 | 数据来源：网站访问日志",
        "llm_summary": llm_summary,
        "llm_context": llm_context,
        "llm_sections": {},
    }


def _report_special_path_query(remote) -> dict[str, Any]:
    return {
        "bool": {
            "should": [
                remote._keyword_contains("uri.keyword", "llms.txt"),
                remote._keyword_contains("uri.keyword", "sitemap"),
            ],
            "minimum_should_match": 1,
        }
    }


def _report_unsuccessful_status_query() -> dict[str, Any]:
    return {"range": {"status": {"gte": 300}}}


def _load_report_llm_config() -> dict[str, str]:
    config = {
        "base_url": os.getenv("REPORT_LLM_BASE_URL", "").strip() or os.getenv("ANTHROPIC_BASE_URL", "").strip(),
        "token": os.getenv("REPORT_LLM_TOKEN", "").strip() or os.getenv("ANTHROPIC_AUTH_TOKEN", "").strip(),
        "model": os.getenv("REPORT_LLM_MODEL", "claude-opus-4-6").strip(),
    }
    if REPORT_LOCAL_CONFIG_PATH.exists():
        try:
            payload = json.loads(REPORT_LOCAL_CONFIG_PATH.read_text(encoding="utf-8"))
            if isinstance(payload, dict):
                config["base_url"] = str(payload.get("base_url") or config["base_url"]).strip()
                config["token"] = str(payload.get("token") or config["token"]).strip()
                config["model"] = str(payload.get("model") or config["model"]).strip()
        except Exception:
            pass
    return config


def _summary_prompt_payload(title_label: str, date_from: str, date_to: str, summary: dict[str, Any], stage: dict[str, Any]) -> dict[str, Any]:
    return {
        "site": title_label,
        "period": [date_from, date_to],
        "ai_index": summary.get("ai_index", 0),
        "ai_search": summary.get("ai_search", 0),
        "ai_training": summary.get("ai_training", 0),
        "seo_bot": summary.get("seo_bot", 0),
        "chatgpt_total": stage.get("chatgpt_total", 0),
        "oai_total": stage.get("oai_total", 0),
        "training_total": stage.get("training_total", 0),
        "first_special_hit": stage.get("first_special_hit"),
    }


def generate_llm_summary_from_context(llm_context: dict[str, Any], config_override: dict[str, str] | None = None) -> str:
    fallback = _stage_summary_text(str(llm_context.get("site") or ""), {
        "chatgpt_total": llm_context.get("chatgpt_total", 0),
        "oai_total": llm_context.get("oai_total", 0),
        "training_total": llm_context.get("training_total", 0),
        "first_special_hit": llm_context.get("first_special_hit"),
    })
    config = _load_report_llm_config()
    if config_override:
        for key in ("base_url", "token", "model"):
            if config_override.get(key):
                config[key] = str(config_override[key]).strip()
    if not config.get("base_url") or not config.get("token"):
        return fallback
    body = {
        "model": config["model"],
        "max_tokens": 320,
        "temperature": 0.2,
        "system": "你是中文数据报告助手。只输出一段正式、简洁、业务风格的总结，不要列表，不要标题，不要解释。",
        "messages": [
            {
                "role": "user",
                "content": f"基于以下JSON写1段中文总结，40到90字，说明当前AI可发现性阶段和下一步观察重点：{json.dumps(llm_context, ensure_ascii=False, separators=(',', ':'))}"
            }
        ],
    }
    try:
        req = request.Request(
            config["base_url"].rstrip("/") + "/v1/messages",
            data=json.dumps(body, ensure_ascii=False).encode("utf-8"),
            headers={
                "content-type": "application/json",
                "x-api-key": config["token"],
                "anthropic-version": "2023-06-01",
            },
            method="POST",
        )
        with request.urlopen(req, timeout=40) as resp:
            payload = json.loads(resp.read().decode("utf-8"))
        parts = payload.get("content") or []
        text = " ".join(str(item.get("text") or "").strip() for item in parts if isinstance(item, dict)).strip()
        return text or fallback
    except (URLError, HTTPError, TimeoutError, json.JSONDecodeError, OSError):
        return fallback


def _report_static_exclusion_query(remote) -> dict[str, Any]:
    return {
        "bool": {
            "must": [remote._static_resource_query()],
            "must_not": [_report_special_path_query(remote)],
        }
    }


def _report_category_queries(remote, shopify_scope: bool) -> dict[str, dict[str, Any]]:
    static_exclusion = _report_static_exclusion_query(remote)
    mirror_non_302 = remote._mirror_non_302_query(allow_io_mirror=shopify_scope)
    probe = remote._suspicious_probe_query()
    exclusions = [static_exclusion, mirror_non_302, probe]
    return {
        "ai_search": {"bool": {"must": [remote._ai_search_query()], "must_not": exclusions}},
        "ai_training": {"bool": {"must": [remote._ai_training_query()], "must_not": exclusions}},
        "ai_index": {"bool": {"must": [remote._ai_index_query()], "must_not": exclusions}},
        "seo_bot": {"bool": {"must": [remote._seo_bot_query()], "must_not": exclusions}},
    }


def _report_base_filters(remote, start_utc: str, end_utc: str, host_filters: list[str] | None, shopify_scope: bool) -> list[dict[str, Any]]:
    return remote._dashboard_base_filters(start_utc, end_utc, host_filters, shopify_scope)


def _query_report_summary(remote, index_names: list[str], host_filters: list[str] | None, start_utc: str, end_utc: str) -> dict[str, Any]:
    shopify_scope = remote._is_shopify_scope(index_names)
    queries = _report_category_queries(remote, shopify_scope)
    filters = _report_base_filters(remote, start_utc, end_utc, host_filters, shopify_scope)
    body = {
        "size": 0,
        "track_total_hits": False,
        "query": {"bool": {"filter": filters}},
        "aggs": {
            "categories": {"filters": {"filters": queries}},
            "days": {
                "date_histogram": {
                    "field": "@timestamp",
                    "calendar_interval": "1d",
                    "time_zone": REPORT_TIMEZONE,
                    "min_doc_count": 0,
                    "extended_bounds": {"min": start_utc, "max": remote._end_bound(end_utc)},
                },
                "aggs": {"categories": {"filters": {"filters": queries}}},
            },
        },
    }
    response = remote._post_json(REPORT_REMOTE_PATH, {"params": {"index": remote._query_index_target(index_names), "body": body}})
    aggs = response["rawResponse"]["aggregations"]
    counts = {key: int(value["doc_count"]) for key, value in aggs["categories"]["buckets"].items()}
    days = []
    for bucket in aggs["days"]["buckets"]:
        item = {"date": bucket["key_as_string"][:10]}
        item.update({key: int(value["doc_count"]) for key, value in bucket["categories"]["buckets"].items()})
        days.append(item)
    counts["ai_total"] = counts.get("ai_search", 0) + counts.get("ai_training", 0) + counts.get("ai_index", 0)
    counts["total"] = counts["ai_total"] + counts.get("seo_bot", 0)
    counts["days"] = days
    return counts


def _query_grouped_rankings(remote, index_names: list[str], host_filters: list[str] | None, start_utc: str, end_utc: str, category_query: dict[str, Any], limit: int = 8) -> list[dict[str, Any]]:
    shopify_scope = remote._is_shopify_scope(index_names)
    body = {
        "size": 0,
        "track_total_hits": False,
        "query": {"bool": {"filter": _report_base_filters(remote, start_utc, end_utc, host_filters, shopify_scope) + [category_query]}},
        "aggs": {
            "uas": {
                "terms": {
                    "field": "ua.keyword",
                    "size": max(limit * 8, 40),
                    "order": {"_count": "desc"},
                }
            }
        },
    }
    response = remote._post_json(REPORT_REMOTE_PATH, {"params": {"index": remote._query_index_target(index_names), "body": body}})
    grouped: dict[str, int] = defaultdict(int)
    for bucket in response["rawResponse"]["aggregations"]["uas"]["buckets"]:
        name = remote.infer_bot_name_from_ua(bucket["key"]) if hasattr(remote, "infer_bot_name_from_ua") else None
        if not name:
            from .bot_taxonomy import infer_bot_name_from_ua
            name = infer_bot_name_from_ua(bucket["key"])
        grouped[name] += int(bucket["doc_count"])
    rows = [{"name": key, "requests": value} for key, value in grouped.items() if key]
    rows.sort(key=lambda item: (-item["requests"], item["name"]))
    return rows[:limit]


def _query_category_rankings(remote, index_names: list[str], host_filters: list[str] | None, start_utc: str, end_utc: str) -> dict[str, list[dict[str, Any]]]:
    shopify_scope = remote._is_shopify_scope(index_names)
    queries = _report_category_queries(remote, shopify_scope)
    return {
        "ai_search": _query_grouped_rankings(remote, index_names, host_filters, start_utc, end_utc, queries["ai_search"], limit=5),
        "ai_training": _query_grouped_rankings(remote, index_names, host_filters, start_utc, end_utc, queries["ai_training"], limit=5),
        "ai_index": _query_grouped_rankings(remote, index_names, host_filters, start_utc, end_utc, queries["ai_index"], limit=5),
        "seo_bot": _query_grouped_rankings(remote, index_names, host_filters, start_utc, end_utc, queries["seo_bot"], limit=5),
    }


def _query_stage_assessment(remote, index_names: list[str], host_filters: list[str] | None, start_utc: str, end_utc: str) -> dict[str, Any]:
    shopify_scope = remote._is_shopify_scope(index_names)
    queries = _report_category_queries(remote, shopify_scope)
    base_filters = _report_base_filters(remote, start_utc, end_utc, host_filters, shopify_scope)
    special_paths = _report_special_path_query(remote)

    first_hit_body = {
        "size": 1,
        "sort": [{"@timestamp": "asc"}],
        "_source": ["@timestamp", "uri", "ua"],
        "query": {"bool": {"filter": base_filters + [special_paths]}},
    }
    first_hit = None
    response = remote._post_json(REPORT_REMOTE_PATH, {"params": {"index": remote._query_index_target(index_names), "body": first_hit_body}})
    hits = response["rawResponse"]["hits"]["hits"]
    if hits:
        source = hits[0].get("_source") or {}
        from .bot_taxonomy import infer_bot_name_from_ua
        first_hit = {
            "date": str(source.get("@timestamp") or "")[:10],
            "uri": source.get("uri") or "",
            "bot": infer_bot_name_from_ua(source.get("ua") or ""),
        }

    training_rank = _query_grouped_rankings(remote, index_names, host_filters, start_utc, end_utc, queries["ai_training"], limit=4)
    training_total = sum(item["requests"] for item in training_rank)

    oai_body = {
        "size": 0,
        "track_total_hits": False,
        "runtime_mappings": {"normalized_page": remote._normalized_page_runtime()},
        "query": {"bool": {"filter": base_filters + [queries["ai_index"], remote._ua_match("oai-searchbot")]}},
        "aggs": {
            "pages": {"cardinality": {"field": "normalized_page"}},
        },
    }
    oai_resp = remote._post_json(REPORT_REMOTE_PATH, {"params": {"index": remote._query_index_target(index_names), "body": oai_body}})
    oai_total = int(oai_resp["rawResponse"]["aggregations"]["pages"]["meta"]["total"] if False else 0)
    oai_pages = int(oai_resp["rawResponse"]["aggregations"]["pages"]["value"] or 0)

    chatgpt_body = {
        "size": 1,
        "sort": [{"@timestamp": "asc"}],
        "_source": ["@timestamp", "uri", "ua"],
        "query": {"bool": {"filter": base_filters + [queries["ai_search"], remote._ua_match("chatgpt-user")]}},
    }
    chatgpt_resp = remote._post_json(REPORT_REMOTE_PATH, {"params": {"index": remote._query_index_target(index_names), "body": chatgpt_body}})
    chatgpt_hits = chatgpt_resp["rawResponse"]["hits"]["hits"]
    chatgpt_total = int(chatgpt_resp["rawResponse"]["hits"]["total"]["value"] if isinstance(chatgpt_resp["rawResponse"]["hits"]["total"], dict) else len(chatgpt_hits))
    chatgpt_first = None
    if chatgpt_hits:
        source = chatgpt_hits[0].get("_source") or {}
        chatgpt_first = {"date": str(source.get("@timestamp") or "")[:10], "uri": source.get("uri") or ""}

    return {
        "first_special_hit": first_hit,
        "training_rank": training_rank,
        "training_total": training_total,
        "oai_total": oai_total,
        "oai_pages": oai_pages,
        "chatgpt_total": chatgpt_total,
        "chatgpt_first": chatgpt_first,
    }


def _query_report_daily(remote, index_names: list[str], host_filters: list[str] | None, start_utc: str, end_utc: str) -> list[dict[str, Any]]:
    return _query_report_summary(remote, index_names, host_filters, start_utc, end_utc)["days"]


def _week_key(day_text: str) -> tuple[date, date]:
    current = date.fromisoformat(day_text)
    start = current - timedelta(days=current.weekday())
    end = start + timedelta(days=6)
    return start, end


def _compute_weeks_in_range(days: list[dict[str, Any]]) -> list[WeekStats]:
    buckets: dict[tuple[date, date], list[dict[str, Any]]] = defaultdict(list)
    for row in days:
        if row.get("ai_search", 0) == 0 and row.get("ai_training", 0) == 0 and row.get("ai_index", 0) == 0 and row.get("seo_bot", 0) == 0:
            continue
        buckets[_week_key(row["date"])].append(row)
    if not buckets:
        today = date.today().isoformat()
        return [WeekStats("W1", today, today, 1, 0, 0, 0, 0)]
    keys = sorted(buckets.keys())
    stats: list[WeekStats] = []
    for idx, key in enumerate(keys, start=1):
        rows = sorted(buckets[key], key=lambda item: item["date"])
        stats.append(
            WeekStats(
                label=f"W{idx}",
                start=rows[0]["date"],
                end=rows[-1]["date"],
                days=len(rows),
                ai_index=sum(int(item.get("ai_index", 0)) for item in rows),
                ai_search=sum(int(item.get("ai_search", 0)) for item in rows),
                ai_training=sum(int(item.get("ai_training", 0)) for item in rows),
                seo_bot=sum(int(item.get("seo_bot", 0)) for item in rows),
            )
        )
    return stats


def _query_top_pages(remote, index_names: list[str], host_filters: list[str] | None, start_utc: str, end_utc: str, limit: int = 25) -> list[dict[str, Any]]:
    shopify_scope = remote._is_shopify_scope(index_names)
    queries = _report_category_queries(remote, shopify_scope)
    focused = remote._dashboard_category_filters(allow_io_mirror=shopify_scope)
    ai_any = {"bool": {"should": [focused["ai_search"], focused["ai_training"], focused["ai_index"]], "minimum_should_match": 1}}
    body = {
        "size": 0,
        "track_total_hits": False,
        "runtime_mappings": {"normalized_page": remote._normalized_page_runtime()},
        "query": {
            "bool": {
                "filter": _report_base_filters(remote, start_utc, end_utc, host_filters, shopify_scope) + [ai_any],
                "must_not": [_report_special_path_query(remote), _report_unsuccessful_status_query()],
            }
        },
        "aggs": {
            "pages": {
                "terms": {"field": "normalized_page", "size": limit},
                "aggs": {
                    "ai_search": {"filter": queries["ai_search"]},
                    "ai_training": {"filter": queries["ai_training"]},
                    "ai_index": {"filter": queries["ai_index"]},
                    "seo_bot": {"filter": queries["seo_bot"]},
                },
            }
        },
    }
    response = remote._post_json(REPORT_REMOTE_PATH, {"params": {"index": remote._query_index_target(index_names), "body": body}})
    rows: list[dict[str, Any]] = []
    for bucket in response["rawResponse"]["aggregations"]["pages"]["buckets"]:
        if str(bucket.get("key") or "") == "/404":
            continue
        row = {
            "page": bucket["key"],
            "ai_search": int(bucket["ai_search"]["doc_count"]),
            "ai_training": int(bucket["ai_training"]["doc_count"]),
            "ai_index": int(bucket["ai_index"]["doc_count"]),
            "seo_bot": int(bucket["seo_bot"]["doc_count"]),
        }
        row["total"] = row["ai_search"] + row["ai_training"] + row["ai_index"] + row["seo_bot"]
        rows.append(row)
    rows.sort(key=lambda item: (-item["total"], item["page"]))
    return rows[:limit]


def _query_stage_assessment_v2(remote, index_names: list[str], host_filters: list[str] | None, start_utc: str, end_utc: str) -> dict[str, Any]:
    shopify_scope = remote._is_shopify_scope(index_names)
    queries = _report_category_queries(remote, shopify_scope)
    base_filters = _report_base_filters(remote, start_utc, end_utc, host_filters, shopify_scope)
    special_paths = _report_special_path_query(remote)

    first_hit_body = {
        "size": 1,
        "sort": [{"@timestamp": "asc"}],
        "_source": ["@timestamp", "uri", "ua"],
        "query": {"bool": {"filter": base_filters + [special_paths]}},
    }
    first_hit = None
    response = remote._post_json(REPORT_REMOTE_PATH, {"params": {"index": remote._query_index_target(index_names), "body": first_hit_body}})
    hits = response["rawResponse"]["hits"]["hits"]
    if hits:
        source = hits[0].get("_source") or {}
        from .bot_taxonomy import infer_bot_name_from_ua
        first_hit = {
            "date": str(source.get("@timestamp") or "")[:10],
            "uri": source.get("uri") or "",
            "bot": infer_bot_name_from_ua(source.get("ua") or ""),
        }

    training_rank = _query_grouped_rankings(remote, index_names, host_filters, start_utc, end_utc, queries["ai_training"], limit=4)
    training_total = sum(item["requests"] for item in training_rank)

    oai_body = {
        "size": 0,
        "track_total_hits": False,
        "runtime_mappings": {"normalized_page": remote._normalized_page_runtime()},
        "query": {"bool": {"filter": base_filters}},
        "aggs": {
            "oai_total": {"filter": {"bool": {"filter": [queries["ai_index"], remote._ua_match("oai-searchbot")]}}},
            "pages": {
                "filter": {"bool": {"filter": [queries["ai_index"], remote._ua_match("oai-searchbot")]}} ,
                "aggs": {"count": {"cardinality": {"field": "normalized_page"}}},
            },
        },
    }
    oai_resp = remote._post_json(REPORT_REMOTE_PATH, {"params": {"index": remote._query_index_target(index_names), "body": oai_body}})
    oai_total = int(oai_resp["rawResponse"]["aggregations"]["oai_total"]["doc_count"] or 0)
    oai_pages = int(oai_resp["rawResponse"]["aggregations"]["pages"]["count"]["value"] or 0)

    chatgpt_first_body = {
        "size": 1,
        "sort": [{"@timestamp": "asc"}],
        "_source": ["@timestamp", "uri", "ua"],
        "query": {"bool": {"filter": base_filters + [queries["ai_search"], remote._ua_match("chatgpt-user")]}}
    }
    chatgpt_first_resp = remote._post_json(REPORT_REMOTE_PATH, {"params": {"index": remote._query_index_target(index_names), "body": chatgpt_first_body}})
    chatgpt_hits = chatgpt_first_resp["rawResponse"]["hits"]["hits"]
    chatgpt_count_body = {
        "size": 0,
        "track_total_hits": False,
        "query": {"bool": {"filter": base_filters}},
        "aggs": {"chatgpt_total": {"filter": {"bool": {"filter": [queries["ai_search"], remote._ua_match("chatgpt-user")]}}}},
    }
    chatgpt_count_resp = remote._post_json(REPORT_REMOTE_PATH, {"params": {"index": remote._query_index_target(index_names), "body": chatgpt_count_body}})
    chatgpt_total = int(chatgpt_count_resp["rawResponse"]["aggregations"]["chatgpt_total"]["doc_count"] or 0)
    chatgpt_first = None
    if chatgpt_hits:
        source = chatgpt_hits[0].get("_source") or {}
        chatgpt_first = {"date": str(source.get("@timestamp") or "")[:10], "uri": source.get("uri") or ""}

    return {
        "first_special_hit": first_hit,
        "training_rank": training_rank,
        "training_total": training_total,
        "oai_total": oai_total,
        "oai_pages": oai_pages,
        "chatgpt_total": chatgpt_total,
        "chatgpt_first": chatgpt_first,
    }


def _inclusive_days(start_text: str, end_text: str) -> int:
    return (date.fromisoformat(end_text) - date.fromisoformat(start_text)).days + 1


def _period_week_note(date_from: str, date_to: str) -> str:
    current = date.fromisoformat(date_from)
    end = date.fromisoformat(date_to)
    chunks = []
    counter = 1
    while current <= end:
        week_end = min(end, current + timedelta(days=(6 - current.weekday())))
        chunks.append(f"W{counter} = {current.month}/{current.day}（周{_weekday_cn(current)}）–{week_end.month}/{week_end.day}（周{_weekday_cn(week_end)}），{(week_end-current).days + 1}天")
        current = week_end + timedelta(days=1)
        counter += 1
    return "周期定义按自然周划分：" + "；".join(chunks) + "。"


def _weekday_cn(day_value: date) -> str:
    return "一二三四五六日"[day_value.weekday()]


def _fill_docx_template(doc, data: dict[str, Any]) -> None:
    paragraphs = doc.paragraphs
    _set_paragraph_text(paragraphs[0], "AI 可发现性流量分析报告")
    _set_paragraph_text(paragraphs[1], f"{data['scope_label']} | {data['month_label']}")
    _set_paragraph_text(paragraphs[2], f"数据周期：{data['date_from']} 至 {data['date_to']}（{data['period_days']} 天）")
    if doc.sections and doc.sections[0].header.paragraphs:
        _set_paragraph_text(doc.sections[0].header.paragraphs[0], f"AI 可发现性流量分析报告  |  Deeplumen  |  {data['date_to']}")

    summary = data["summary"]
    stage = data["stage"]
    weeks = data["week_stats"]
    rankings = data["rankings"]
    if data["is_shopify"]:
        _fill_shopify_template(doc, data, summary, stage, weeks, rankings)
    else:
        _fill_non_shopify_template(doc, data, summary, stage, weeks, rankings)


def _fill_shopify_template(doc, data: dict[str, Any], summary: dict[str, Any], stage: dict[str, Any], weeks: list[WeekStats], rankings: dict[str, list[dict[str, Any]]]) -> None:
    paragraphs = doc.paragraphs
    llm = data.get("llm_sections", {}) or {}
    llm_mode = _has_llm_findings(llm)
    _set_paragraph_text(paragraphs[14], "二、核心发现" if llm_mode else "二、核心观察")
    _set_paragraph_text(paragraphs[15], "以下是本周期最重要的三项发现，按照商业意义排序：" if llm_mode else "以下为本周期最值得关注的三项观察，按业务影响排序：")
    _set_paragraph_text(paragraphs[17], _narrative_heading(_shopify_finding_title_1(data, stage), llm_mode))
    _set_paragraph_text(paragraphs[18], llm.get("finding1_body") or _shopify_finding_body_1(data, stage))
    _set_paragraph_text(paragraphs[20], _narrative_heading(_shopify_finding_title_2(data, stage), llm_mode))
    _set_paragraph_text(paragraphs[21], llm.get("finding2_body") or _shopify_finding_body_2(data, stage))
    _set_paragraph_text(paragraphs[23], _narrative_heading(_shopify_finding_title_3(data), llm_mode))
    _set_paragraph_text(paragraphs[24], llm.get("finding3_body") or _shopify_finding_body_3(data))
    _fill_anomaly_section(doc, data, shopify=True)
    _set_paragraph_text(paragraphs[39], "4.1  AI 可发现性阶段现状")
    _set_paragraph_text(paragraphs[40], f"▍ {data['title_label']} 在 {data['period_days']} 天内完成四个阶段，目前处于第四阶段（推荐展示）且增长明显。")
    _set_paragraph_text(paragraphs[43], "4.2  访问量汇总")
    _set_paragraph_text(paragraphs[44], f"▍ {data['period_days']} 天内共 {summary['total']:,} 次 Bot 访问，其中 {(_share(summary['ai_total'], summary['total']))} 为 AI 相关流量。")
    _set_paragraph_text(paragraphs[47], f"统计口径：仅含可识别的 Bot 类流量，不含浏览器访问。数据周期 {data['date_from']} 至 {data['date_to']}（{data['period_days']} 天）。")
    _set_paragraph_text(paragraphs[49], "4.3  每日流量趋势")
    _set_paragraph_text(paragraphs[50], llm.get("daily_summary") or _daily_trend_summary(summary))
    _set_paragraph_text(paragraphs[53], "图注：展示总量、AI 索引、AI 训练、AI 搜索与 SEO 五条每日趋势线。")
    _replace_list_texts(doc, 55, _daily_trend_bullets(summary, stage))
    _set_paragraph_text(paragraphs[60], "4.4  周维度趋势")
    _set_paragraph_text(paragraphs[61], llm.get("weekly_summary") or _shopify_weekly_summary(weeks))
    _replace_list_texts(doc, 66, _weekly_total_bullets(weeks)[:3])
    _set_paragraph_text(paragraphs[69], "4.5  产品页访问分布（Top 5）")
    _set_paragraph_text(paragraphs[70], llm.get("pages_summary") or "▍ 重点商品页的访问已明显集中在少数高频品类。")
    _set_paragraph_text(paragraphs[73], "图注：展示 Top 5 产品页在 AI 索引、AI 搜索、AI 训练与 SEO 之间的访问分布。")
    _replace_list_texts(doc, 75, _top_pages_bullets(data["top_pages"], stage)[:2])
    _set_paragraph_text(paragraphs[78], "4.6  Shopify 店铺数据对比")
    _set_paragraph_text(paragraphs[79], llm.get("compare_summary") or _shopify_store_compare_summary(data))
    _set_paragraph_text(paragraphs[81], "说明：本节对比站点总量、AI 流量与用户流量的变化关系，重点观察增长是否同步。")
    _set_paragraph_text(paragraphs[84], _shopify_store_compare_bullet(data))
    _set_paragraph_text(paragraphs[87], "左轴（柱）为站点总请求量；右轴（线）为 AI 请求量。")
    _set_paragraph_text(paragraphs[90], "柱状为用户与 Bot 的日总量结构，折线为 AI 请求量变化。")
    _set_paragraph_text(paragraphs[92], _shopify_store_compare_note_1(data))
    _set_paragraph_text(paragraphs[93], _shopify_store_compare_note_2(data))
    _fill_stage_table(doc.tables[2], stage)
    _fill_summary_table(doc.tables[3], summary, rankings, data["period_days"], include_seo=True)
    _fill_week_table(doc.tables[4], weeks, include_seo=True)
    _fill_shopify_compare_table(doc.tables[5], weeks, data)
    _fill_target_table(doc.tables[6], stage)


def _fill_non_shopify_template(doc, data: dict[str, Any], summary: dict[str, Any], stage: dict[str, Any], weeks: list[WeekStats], rankings: dict[str, list[dict[str, Any]]]) -> None:
    paragraphs = doc.paragraphs
    llm = data.get("llm_sections", {}) or {}
    _set_paragraph_text(paragraphs[10], "1.2  AI BOT流量观察框架")
    _set_paragraph_text(paragraphs[11], "本报告重点观察 AI BOT流量规模、AI BOT平台构成变化与重点页面覆盖情况。")
    llm_mode = _has_llm_findings(llm)
    _set_paragraph_text(paragraphs[14], "二、核心发现" if llm_mode else "二、核心观察")
    _set_paragraph_text(paragraphs[15], "以下是本周期最重要的三项发现，重点聚焦 AI BOT流量增长和 AI BOT平台结构变化：" if llm_mode else "以下为本周期最值得关注的三项观察，重点聚焦 AI BOT流量增长与 AI BOT平台结构变化：")
    _set_paragraph_text(paragraphs[17], _narrative_heading(_non_shopify_finding_title_1(data), llm_mode))
    _set_paragraph_text(paragraphs[18], llm.get("finding1_body") or _non_shopify_finding_body_1(data))
    _set_paragraph_text(paragraphs[20], _narrative_heading(_non_shopify_finding_title_2(data, weeks), llm_mode))
    _set_paragraph_text(paragraphs[21], llm.get("finding2_body") or _non_shopify_finding_body_2(data, weeks))
    _set_paragraph_text(paragraphs[23], _narrative_heading(_non_shopify_finding_title_3(data), llm_mode))
    _set_paragraph_text(paragraphs[24], llm.get("finding3_body") or _non_shopify_finding_body_3(data))
    _fill_anomaly_section(doc, data, shopify=False)
    _set_paragraph_text(paragraphs[39], "4.1  AI BOT流量现状")
    _set_paragraph_text(paragraphs[40], f"▍ {data['title_label']} 在当前周期内已形成稳定的 AI BOT访问结构，重点关注 AI BOT流量变化与平台构成演进。")
    _set_paragraph_text(paragraphs[43], "4.2  访问量汇总")
    _set_paragraph_text(paragraphs[44], f"▍ {data['period_days']} 天内共 {summary['ai_total']:,} 次 AI BOT流量，AI 索引、AI 搜索、AI 训练三类 AI BOT共同构成当前结构。")
    _set_paragraph_text(paragraphs[47], f"统计口径：仅统计 AI 索引、AI 搜索、AI 训练三类 AI BOT流量；不包含 SEO BOT、用户流量及失败访问。数据周期 {data['date_from']} 至 {data['date_to']}（{data['period_days']} 天）。")
    _set_paragraph_text(paragraphs[49], "4.3  每日流量趋势")
    _set_paragraph_text(paragraphs[50], llm.get("daily_summary") or _daily_trend_summary_non_shopify(summary))
    _set_paragraph_text(paragraphs[53], "图注：展示 AI BOT总量、AI 索引、AI 训练、AI 搜索四条每日趋势线。")
    _replace_list_texts(doc, 55, _daily_trend_bullets_non_shopify(summary, stage))
    _set_paragraph_text(paragraphs[60], "4.4  周维度趋势")
    _set_paragraph_text(paragraphs[61], llm.get("weekly_summary") or _non_shopify_weekly_summary(weeks))
    _replace_list_texts(doc, 66, _weekly_total_bullets(weeks)[:3])
    _set_paragraph_text(paragraphs[69], "4.5  重点页面访问分布（Top 5）")
    _set_paragraph_text(paragraphs[70], llm.get("pages_summary") or "▍ AI BOT访问正在向部分重点页面集中，可据此判断哪些页面最容易被 AI 平台读取与引用。")
    _set_paragraph_text(paragraphs[73], "图注：展示 Top 5 重点页面在 AI 索引、AI 搜索与 AI 训练之间的访问分布。")
    _replace_list_texts(doc, 75, _non_shopify_top_pages_bullets(data["top_pages"], stage))
    _set_paragraph_text(paragraphs[78], "4.6  AI BOT平台构成变化")
    _set_paragraph_text(paragraphs[79], llm.get("compare_summary") or _non_shopify_client_flow_summary(data))
    _set_paragraph_text(paragraphs[81], "说明：本节仅观察 AI BOT总量、AI BOT平台构成和周维度平均值变化，不纳入人类流量与 SEO BOT。")
    _set_paragraph_text(paragraphs[84], _non_shopify_client_flow_note_1(data))
    _set_paragraph_text(paragraphs[87], "左轴（柱）为 AI BOT总请求量；右轴（线）为 AI 搜索 BOT请求量。")
    _set_paragraph_text(paragraphs[90], "柱状展示 AI BOT总量，折线展示 AI 搜索 BOT变化，用于观察高价值 AI BOT访问是否同步上升。")
    _set_paragraph_text(paragraphs[92], _non_shopify_client_flow_note_2(data))
    _set_paragraph_text(paragraphs[93], _non_shopify_client_flow_note_3(data))
    _set_paragraph_text(paragraphs[95], "五、后续建议")
    _set_paragraph_text(paragraphs[96], "基于当前数据，建议在以下几个方向持续关注：")
    _set_paragraph_text(paragraphs[98], "5.1  持续跟踪 AI BOT平台占比变化")
    _set_paragraph_text(paragraphs[99], _non_shopify_followup_platform_note(data))
    _set_paragraph_text(paragraphs[102], "5.2  持续跟踪重点页面变化")
    _set_paragraph_text(paragraphs[103], "建议定期（每 2 周）统计重点页面的 AI BOT访问变化，重点关注：")
    _set_paragraph_text(paragraphs[104], _non_shopify_followup_page_note_1(data))
    _set_paragraph_text(paragraphs[105], _non_shopify_followup_page_note_2(data))
    _set_paragraph_text(paragraphs[106], _non_shopify_followup_page_note_3(data))
    _fill_non_shopify_glossary_table(doc.tables[0], data)
    _fill_observation_table(doc.tables[1], data)
    _fill_current_focus_table(doc.tables[2], data, weeks)
    _fill_summary_table(doc.tables[3], summary, rankings, data["period_days"], include_seo=False)
    _fill_week_table(doc.tables[4], weeks, include_seo=False)
    _fill_non_shopify_compare_table(doc.tables[5], weeks, data)
    _fill_recommend_target_table(doc.tables[6], data)


def _set_paragraph_text(paragraph, text: str) -> None:
    paragraph.text = text


def _has_llm_findings(llm_sections: dict[str, Any] | None) -> bool:
    if not isinstance(llm_sections, dict):
        return False
    for key in ("finding1_body", "finding2_body", "finding3_body"):
        value = llm_sections.get(key)
        if isinstance(value, str) and value.strip():
            return True
    return False


def _narrative_heading(text: str, llm_mode: bool) -> str:
    if llm_mode:
        return text
    return re.sub(r"^发现(?=\s*\d|[ 　])", "观察", text, count=1)


def _replace_list_texts(doc, start_index: int, items: list[str]) -> None:
    for offset, text in enumerate(items):
        _set_paragraph_text(doc.paragraphs[start_index + offset], text)


def _stage_summary_text(title_label: str, stage: dict[str, Any]) -> str:
    if stage.get("chatgpt_total", 0) > 0:
        return f"{title_label} 已完整经历四个阶段，目前处于第四阶段。下一步重点观察 ChatGPT-User 点击量是否开始稳定增长。AI发现、收录、建立索引的时间跨度通常在数周左右，期间AI流量表现会相对平稳。"
    if stage.get("oai_total", 0) > 0:
        return f"{title_label} 已进入 AI 索引阶段。当前重点观察 ChatGPT-User 是否开始出现真实点击，以及训练型抓取是否维持稳定。"
    if stage.get("training_total", 0) >= 5:
        return f"{title_label} 已进入训练收录阶段。当前重点观察 OAI-SearchBot 是否开始批量建立索引，以及入口信号是否继续增长。"
    if stage.get("first_special_hit"):
        return f"{title_label} 已进入 AI 发现阶段。当前重点观察训练型爬虫是否开始出现稳定访问，以及后续的索引建立。"
    return f"{title_label} 暂未观察到明确的 AI 发现与索引信号，建议继续观察入口信号与后续索引变化。"


def _daily_trend_summary(summary: dict[str, Any]) -> str:
    days = summary["days"]
    if not days:
        return "当前周期暂无可用数据。"
    peak = max(days, key=lambda item: item.get("ai_search", 0) + item.get("ai_training", 0) + item.get("ai_index", 0) + item.get("seo_bot", 0))
    peak_total = peak.get("ai_search", 0) + peak.get("ai_training", 0) + peak.get("ai_index", 0) + peak.get("seo_bot", 0)
    return f"{len(days)}天内流量先出现阶段性峰值，随后进入相对稳定阶段。峰值出现在 {peak['date']}，单日 {peak_total:,} 次。"


def _daily_trend_bullets(summary: dict[str, Any], stage: dict[str, Any]) -> list[str]:
    days = summary["days"]
    if not days:
        return ["当前周期暂无趋势数据。", "等待更多流量后再观察阶段性变化。", "—", "—"]
    peak = max(days, key=lambda item: item.get("ai_search", 0) + item.get("ai_training", 0) + item.get("ai_index", 0) + item.get("seo_bot", 0))
    peak_total = peak.get("ai_search", 0) + peak.get("ai_training", 0) + peak.get("ai_index", 0) + peak.get("seo_bot", 0)
    non_peak = [row for row in days if row["date"] != peak["date"]]
    stable_avg = round(sum((row.get("ai_search", 0) + row.get("ai_training", 0) + row.get("ai_index", 0) + row.get("seo_bot", 0)) for row in non_peak) / len(non_peak), 1) if non_peak else peak_total
    bullets = [
        f"{peak['date']}：单日 {peak_total:,} 次，为观测期峰值。",
        f"其余日期进入相对稳定阶段，日均约 {stable_avg} 次。",
    ]
    if stage.get("chatgpt_total", 0) > 0 and stage.get("chatgpt_first"):
        bullets.append(f"{stage['chatgpt_first']['date']}：ChatGPT-User 首次出现，共 {stage['chatgpt_total']} 次。")
    else:
        bullets.append("当前周期内暂未观察到 ChatGPT-User 真实点击。")
    bullets.append(f"AI 训练流量累计 {summary.get('ai_training', 0):,} 次，保持稳定采集。")
    return bullets


def _weekly_total_bullets(weeks: list[WeekStats]) -> list[str]:
    if not weeks:
        return ["当前时间范围内暂无周维度数据。", "—", "—", "—"]
    first = weeks[0]
    last = weeks[-1]
    change = _pct_change(last.daily_avg(last.total), first.daily_avg(first.total)) if len(weeks) > 1 else "N/A"
    peak = max(weeks, key=lambda item: item.daily_avg(item.total))
    return [
        f"首周（{_short_range(first)}，{first.days}天）：{first.total:,} 次，日均 {first.daily_avg(first.total):.1f} 次。",
        f"末周（{_short_range(last)}，{last.days}天）：{last.total:,} 次，日均 {last.daily_avg(last.total):.1f} 次。",
        f"首末周日均变化 {change}，峰值出现在 {peak.label}（{_short_range(peak)}）。",
        f"当前时间范围内共纳入 {len(weeks)} 个自然周进行对比。"
    ]


def _weekly_ai_bullets(weeks: list[WeekStats]) -> list[str]:
    if not weeks:
        return ["当前时间范围内暂无 AI 周维度数据。", "—", "—"]
    peak_index = max(weeks, key=lambda item: item.daily_avg(item.ai_index))
    peak_training = max(weeks, key=lambda item: item.daily_avg(item.ai_training))
    peak_search = max(weeks, key=lambda item: item.daily_avg(item.ai_search))
    return [
        f"AI索引峰值周：{peak_index.label}（{_short_range(peak_index)}），日均 {peak_index.daily_avg(peak_index.ai_index):.1f} 次。",
        f"AI训练峰值周：{peak_training.label}（{_short_range(peak_training)}），日均 {peak_training.daily_avg(peak_training.ai_training):.1f} 次。",
        f"AI搜索峰值周：{peak_search.label}（{_short_range(peak_search)}），日均 {peak_search.daily_avg(peak_search.ai_search):.1f} 次。"
    ]


def _weekly_seo_bullets(seo_rankings: list[dict[str, Any]], weeks: list[WeekStats]) -> list[str]:
    names = [item["name"] for item in seo_rankings[:3]] if seo_rankings else ["SEO Bot"]
    bullets = []
    for name in names[:3]:
        bullets.append(f"{name}：用于观察当前时间范围内各自然周的 SEO 工具抓取变化。")
    while len(bullets) < 3:
        bullets.append("当前周期内该类 SEO Bot 数据较少。")
    return bullets


def _top_pages_bullets(top_pages: list[dict[str, Any]], stage: dict[str, Any]) -> list[str]:
    if not top_pages:
        return ["当前周期暂无产品页 Bot 访问数据。", "—", "—"]
    top_three = "、".join(item["page"] for item in top_pages[:3])
    bullets = [
        f"高频页面集中在：{top_three}。",
        f"OAI-SearchBot 累计访问 {stage.get('oai_total', 0):,} 次，覆盖约 {stage.get('oai_pages', 0):,} 个独立页面。",
        "Top 页面通常会同时被 AI 索引、AI 训练和 SEO Bot 重复访问，可作为重点优化对象。"
    ]
    return bullets


def _fill_stage_table(table, stage: dict[str, Any]) -> None:
    rows = [
        ("第一阶段", "发现", "入口信号首次出现", _stage_one_status(stage)),
        ("第二阶段", "训练收录", "AI训练爬虫≥5次", _stage_two_status(stage)),
        ("第三阶段", "AI索引", "OAI-SearchBot出现", _stage_three_status(stage)),
        ("第四阶段 ★", "推荐展示", "ChatGPT-User点击", _stage_four_status(stage)),
    ]
    _clear_table(table)
    for idx, row_values in enumerate(rows, start=1):
        for col, value in enumerate(row_values):
            table.cell(idx, col).text = value


def _stage_one_status(stage: dict[str, Any]) -> str:
    hit = stage.get("first_special_hit")
    if not hit:
        return "待确认"
    return f"✅ 已完成（{hit['date']} {hit['bot']} 首次读取）"


def _stage_two_status(stage: dict[str, Any]) -> str:
    if stage.get("training_total", 0) < 5:
        return "待确认"
    details = " + ".join(f"{item['name']} {item['requests']}次" for item in stage.get("training_rank", [])[:4])
    return f"✅ 已完成（{details}）"


def _stage_three_status(stage: dict[str, Any]) -> str:
    if stage.get("oai_total", 0) <= 0:
        return "待确认"
    return f"✅ 已完成（{stage['oai_total']:,}次，覆盖{stage['oai_pages']:,}个页面）"


def _stage_four_status(stage: dict[str, Any]) -> str:
    if stage.get("chatgpt_total", 0) <= 0:
        return "待确认"
    return f"★ 已确认（{stage['chatgpt_first']['date']} ChatGPT-User {stage['chatgpt_total']}次）"


def _fill_summary_table(table, summary: dict[str, Any], rankings: dict[str, list[dict[str, Any]]], period_days: int, include_seo: bool = True) -> None:
    share_total = summary.get("total", 0) if include_seo else summary.get("ai_total", 0)
    rows = [
        ("来源类型", "代表平台", "累计访问次数", "日均", "占比", "说明"),
        ("AI索引", _top_names(rankings.get("ai_index")), summary.get("ai_index", 0), _daily_avg(summary.get("ai_index", 0), period_days), _share(summary.get("ai_index", 0), share_total), "AI 索引 BOT流量"),
        ("AI搜索 ★", _top_names(rankings.get("ai_search")), summary.get("ai_search", 0), _daily_avg(summary.get("ai_search", 0), period_days), _share(summary.get("ai_search", 0), share_total), "AI 搜索 BOT流量"),
        ("AI训练", _top_names(rankings.get("ai_training")), summary.get("ai_training", 0), _daily_avg(summary.get("ai_training", 0), period_days), _share(summary.get("ai_training", 0), share_total), "AI 训练 BOT流量"),
    ]
    if include_seo:
        rows.append(("AI合计", "—", summary.get("ai_total", 0), _daily_avg(summary.get("ai_total", 0), period_days), _share(summary.get("ai_total", 0), summary.get("total", 0)), "AI索引 + 搜索 + 训练"))
        rows.append(("SEO Bot", _top_names(rankings.get("seo_bot")), summary.get("seo_bot", 0), _daily_avg(summary.get("seo_bot", 0), period_days), _share(summary.get("seo_bot", 0), summary.get("total", 0)), "传统搜索与SEO生态"))
        rows.append(("总计", "—", summary.get("total", 0), _daily_avg(summary.get("total", 0), period_days), "100%", ""))
    else:
        rows.append(("总计", "—", summary.get("ai_total", 0), _daily_avg(summary.get("ai_total", 0), period_days), "100%", "AI BOT总量"))
    while len(table.rows) < len(rows):
        table.add_row()
    _clear_table(table)
    for idx, row in enumerate(rows):
        for col, value in enumerate(row):
            table.cell(idx, col).text = str(value)
    _trim_table_rows(table, len(rows))


def _fill_taxonomy_table(table, rows: list[tuple[str, str]]) -> None:
    current = len(table.rows)
    target = len(rows) + 1
    for _ in range(target - current):
        table.add_row()
    for idx, (bot_name, category) in enumerate(rows, start=1):
        table.cell(idx, 0).text = bot_name
        table.cell(idx, 1).text = category


def _fill_non_shopify_glossary_table(table, data: dict[str, Any]) -> None:
    search_names = _top_names(data.get("rankings", {}).get("ai_search")).replace(" / ", "、")
    training_names = _top_names(data.get("rankings", {}).get("ai_training")).replace(" / ", "、")
    index_names = _top_names(data.get("rankings", {}).get("ai_index")).replace(" / ", "、")
    rows = [
        ("名词", "含义"),
        ("AI 训练爬虫", f"用于采集网页内容并纳入模型训练的数据抓取程序。本周期主要贡献平台包括：{training_names or '—'}。"),
        ("AI 索引爬虫", f"用于建立 AI 搜索候选库的抓取程序。本周期主要索引平台包括：{index_names or '—'}。"),
        ("AI 搜索 BOT", f"代表具体 AI 平台在回答或推荐过程中发起的访问。本周期主要来源包括：{search_names or '—'}。"),
        ("重点页面", "指在当前周期内被 AI 平台高频访问、最值得继续分析与优化的页面。"),
        ("周维度日均", "按自然周计算的 AI BOT日均访问量，用于观察增长趋势与平台结构变化。"),
    ]
    while len(table.rows) < len(rows):
        table.add_row()
    _clear_table(table)
    for r, row in enumerate(rows):
        for c, value in enumerate(row):
            table.cell(r, c).text = value
    _trim_table_rows(table, len(rows))


def _shopify_finding_title_1(data: dict[str, Any], stage: dict[str, Any]) -> str:
    return f"发现 1　ChatGPT 已开始向真实用户推荐 {data['title_label']} 页面"


def _shopify_finding_body_1(data: dict[str, Any], stage: dict[str, Any]) -> str:
    total = stage.get("chatgpt_total", 0)
    first_date = (stage.get("chatgpt_first") or {}).get("date") or data["date_from"]
    return f"{first_date} 起，ChatGPT-User 已开始产生真实访问，本周期累计 {total} 次，说明 AI 推荐链路已经建立，并开始向用户展示站点内容。"


def _shopify_finding_title_2(data: dict[str, Any], stage: dict[str, Any]) -> str:
    return f"发现 2　OpenAI 已完成对 {data['title_label']} 目录的系统性建索引"


def _shopify_finding_body_2(data: dict[str, Any], stage: dict[str, Any]) -> str:
    return f"当前 OAI-SearchBot 累计访问 {stage.get('oai_total', 0):,} 次，覆盖约 {stage.get('oai_pages', 0):,} 个页面，说明页面已进入 AI 搜索候选范围。"


def _shopify_finding_title_3(data: dict[str, Any]) -> str:
    return "发现 3　AI 推荐带来的真实访客仍处于早期阶段"


def _shopify_finding_body_3(data: dict[str, Any]) -> str:
    dashboard = data["dashboard_view"]
    return f"当前总访问 {dashboard['cards']['total_requests']:,} 次，其中 AI 流量 {dashboard['cards']['ai_requests']:,} 次。AI 推荐链路已经形成，但要持续观察其是否稳定转化为真实访问增长。"


def _non_shopify_finding_title_1(data: dict[str, Any]) -> str:
    return f"发现 1　{data['title_label']} 的 AI BOT流量已形成稳定规模"


def _non_shopify_finding_body_1(data: dict[str, Any]) -> str:
    summary = data["summary"]
    search_names = _top_names(data.get("rankings", {}).get("ai_search")).replace(" / ", "、")
    return f"本周期 AI BOT流量共 {summary['ai_total']:,} 次，其中 AI 搜索 BOT {summary['ai_search']:,} 次、AI 训练 BOT {summary['ai_training']:,} 次、AI 索引 BOT {summary['ai_index']:,} 次；搜索侧主要由 {search_names} 拉动。"


def _non_shopify_finding_title_2(data: dict[str, Any], weeks: list[WeekStats]) -> str:
    return "发现 2　AI BOT访问占比最高的平台已经非常清晰"


def _non_shopify_finding_body_2(data: dict[str, Any], weeks: list[WeekStats]) -> str:
    if not weeks:
        return "当前时间范围内暂无足够周数据判断 AI BOT平台构成变化。"
    top_platform = _overall_top_ai_platform(data.get("rankings", {}))
    top_name = str(top_platform.get("name") or "—")
    top_requests = int(top_platform.get("requests") or 0)
    share = _share(top_requests, data["summary"].get("ai_total", 0))
    return f"本周期 AI BOT访问占比最高的平台是 {top_name}，累计 {top_requests:,} 次，占全部 AI BOT流量的 {share}。"


def _non_shopify_finding_title_3(data: dict[str, Any]) -> str:
    return "发现 3　重点页面已形成清晰的 AI 抓取焦点"


def _non_shopify_finding_body_3(data: dict[str, Any]) -> str:
    top_pages = data["top_pages"]
    if not top_pages:
        return "当前周期内暂无明显的页面访问集中现象。"
    top = top_pages[0]
    names = "、".join(item["page"] for item in top_pages[:5])
    return f"当前 AI BOT访问最高的页面是 {top['page']}，累计 {top['total']:,} 次；其后依次集中在 {names} 等页面，建议优先围绕这些页面展开内容优化。"


def _shopify_weekly_summary(weeks: list[WeekStats]) -> str:
    if not weeks:
        return "当前时间范围内暂无可用周数据。"
    return f"▍ 当前时间范围内共覆盖 {len(weeks)} 个自然周，AI 流量结构已从“被索引”逐步转向“被推荐”，建议重点观察末周 ChatGPT-User 的持续性。"


def _daily_trend_summary_non_shopify(summary: dict[str, Any]) -> str:
    days = summary["days"]
    if not days:
        return "当前周期暂无可用数据。"
    peak = max(days, key=lambda item: item.get("ai_search", 0) + item.get("ai_training", 0) + item.get("ai_index", 0))
    peak_total = peak.get("ai_search", 0) + peak.get("ai_training", 0) + peak.get("ai_index", 0)
    return f"▍ 当前时间范围内 AI BOT流量总体呈波动变化，峰值出现在 {peak['date']}，单日 {peak_total:,} 次；后续可结合 AI 搜索 BOT与训练 BOT的分项变化判断增长来源。"


def _daily_trend_bullets_non_shopify(summary: dict[str, Any], stage: dict[str, Any]) -> list[str]:
    days = summary["days"]
    if not days:
        return ["当前周期暂无趋势数据。", "—", "—", "—"]
    peak = max(days, key=lambda item: item.get("ai_search", 0) + item.get("ai_training", 0) + item.get("ai_index", 0))
    return [
        f"{peak['date']}：AI BOT总量峰值日，单日 {peak.get('ai_search', 0) + peak.get('ai_training', 0) + peak.get('ai_index', 0):,} 次。",
        f"AI 索引 BOT累计 {summary.get('ai_index', 0):,} 次，AI 搜索 BOT累计 {summary.get('ai_search', 0):,} 次，AI 训练 BOT累计 {summary.get('ai_training', 0):,} 次。",
        "可持续跟踪峰值日后的回落速度，以判断增长是否稳定。",
        "本节仅关注 AI BOT流量变化，不纳入 SEO BOT 与用户流量。",
    ]


def _non_shopify_weekly_summary(weeks: list[WeekStats]) -> str:
    if not weeks:
        return "▍ 当前时间范围内暂无周维度数据。"
    first, last = weeks[0], weeks[-1]
    return f"▍ 当前时间范围内共覆盖 {len(weeks)} 个自然周，首周 AI BOT日均 {first.daily_avg(first.ai_total):.1f} 次，末周 AI BOT日均 {last.daily_avg(last.ai_total):.1f} 次，可用于观察 AI BOT平均值与平台构成的变化趋势。"


def _non_shopify_top_pages_bullets(top_pages: list[dict[str, Any]], stage: dict[str, Any]) -> list[str]:
    if not top_pages:
        return ["当前周期暂无重点页面访问数据。", "—"]
    top_three = "、".join(f"{item['page']}（{item['total']:,}次）" for item in top_pages[:3])
    return [
        f"重点页面主要集中在：{top_three}。",
        "建议优先优化这些高频页面的标题、结构与内容表达，并结合具体页面主题判断哪些内容最容易被 AI 平台抓取与推荐。"
    ]


def _shopify_store_compare_summary(data: dict[str, Any]) -> str:
    dashboard = data["dashboard_view"]
    return f"▍ 当前周期内总访问 {dashboard['cards']['total_requests']:,} 次，AI 流量 {dashboard['cards']['ai_requests']:,} 次。Shopify 站点适合继续观察 AI 曝光是否向真实推荐转化。"


def _shopify_store_compare_bullet(data: dict[str, Any]) -> str:
    return "如 Shopify 后台与 Agentic Page 数据出现同步异常放大，应优先排查投流或搜索引擎批量抓取事件。"


def _shopify_store_compare_note_1(data: dict[str, Any]) -> str:
    return "上线后的 AI 请求量可与站内会话一起观察，重点判断是否出现稳定的人类访问转化。"


def _shopify_store_compare_note_2(data: dict[str, Any]) -> str:
    return "若 Bot 与 Human 同步放大，通常说明是外部流量事件；若只有 AI 搜索增长，则更接近推荐增长。"


def _non_shopify_client_flow_summary(data: dict[str, Any]) -> str:
    summary = data["summary"]
    return f"▍ 当前周期 AI BOT总量 {summary['ai_total']:,} 次，其中 AI 搜索 BOT {summary['ai_search']:,} 次，是当前最值得持续观察的高价值 AI BOT来源。"


def _non_shopify_client_flow_note_1(data: dict[str, Any]) -> str:
    weeks = data["week_stats"]
    if not weeks:
        return "当前时间范围内暂无足够数据。"
    first, last = weeks[0], weeks[-1]
    return f"首周 AI BOT日均 {first.daily_avg(first.ai_total):.1f} 次，末周 AI BOT日均 {last.daily_avg(last.ai_total):.1f} 次，可直接观察 AI BOT平均值是否在增长。"


def _non_shopify_client_flow_note_2(data: dict[str, Any]) -> str:
    weeks = data["week_stats"]
    if not weeks:
        return "暂无周维度 AI 组成数据。"
    first, last = weeks[0], weeks[-1]
    return f"首周 AI BOT日均 {first.daily_avg(first.ai_total):.1f} 次，末周 AI BOT日均 {last.daily_avg(last.ai_total):.1f} 次；AI BOT平台构成变化是核心观察指标。"


def _non_shopify_client_flow_note_3(data: dict[str, Any]) -> str:
    return "建议持续结合周维度变化，判断 AI BOT总量增长与平台构成演进是否同步。"


def _non_shopify_followup_platform_note(data: dict[str, Any]) -> str:
    top_platform = _overall_top_ai_platform(data.get("rankings", {}))
    top_name = str(top_platform.get("name") or "—")
    top_requests = int(top_platform.get("requests") or 0)
    share = _share(top_requests, data["summary"].get("ai_total", 0))
    return f"当前占比最高的 AI BOT平台是 {top_name}，累计 {top_requests:,} 次，占全部 AI BOT流量的 {share}。建议持续观察它的占比变化，以及是否有新的平台进入前列。"


def _non_shopify_followup_page_note_1(data: dict[str, Any]) -> str:
    top_pages = data.get("top_pages") or []
    if not top_pages:
        return "① 当前周期内暂无稳定的重点页面，可继续观察后续 AI BOT访问是否开始集中。"
    top_three = "、".join(f"{item['page']}（{item['total']:,}次）" for item in top_pages[:3])
    return f"① 当前访问量 Top 3 的重点页面为 {top_three}，可优先作为内容优化对象。"


def _non_shopify_followup_page_note_2(data: dict[str, Any]) -> str:
    return "② 关注新进入前列的重点页面，判断 AI BOT关注主题是否正在变化。"


def _non_shopify_followup_page_note_3(data: dict[str, Any]) -> str:
    return "③ 对长期没有 AI BOT访问的重点栏目，检查页面入口、内容结构和可读性是否存在问题。"


def _fill_observation_table(table, data: dict[str, Any]) -> None:
    rows = [
        ("观察维度", "当前关注点", "关键指标", "说明"),
        ("AI BOT总量", "AI BOT总访问与周均变化", f"{data['summary']['ai_total']:,} 次", "用于判断 AI BOT流量是否稳定增长。"),
        ("AI BOT构成", "AI 索引 / AI 搜索 / AI 训练", f"{data['summary']['ai_index']:,} / {data['summary']['ai_search']:,} / {data['summary']['ai_training']:,}", "用于观察 AI BOT平台结构是否发生变化。"),
        ("重点页面覆盖", "Top 5 页面分布", "已生成", "用于判断 AI 访问正在集中到哪些重点页面。"),
        ("异常波动", "显著峰值与结构变化", "见第三章", "仅总结有明确数据支撑的异常现象。"),
    ]
    while len(table.rows) < len(rows):
        table.add_row()
    _clear_table(table)
    for r, row in enumerate(rows):
        for c, value in enumerate(row[: min(len(row), len(table.rows[r].cells))]):
            table.cell(r, c).text = value


def _fill_current_focus_table(table, data: dict[str, Any], weeks: list[WeekStats]) -> None:
    rows = [
        ("观察项", "当前口径", "关键指标", "说明"),
        ("AI BOT总量", "AI BOT总访问变化", f"{data['summary']['ai_total']:,} 次", "关注 AI BOT平均值和总量变化"),
        ("AI BOT构成", "AI 索引 / 搜索 / 训练", f"{data['summary']['ai_index']:,} / {data['summary']['ai_search']:,} / {data['summary']['ai_training']:,}", "关注三类 AI BOT平台比例变化"),
        ("重点页面", "Top 5 页面", "已生成", "观察 AI 访问集中到哪些页面"),
        ("异常", "显著波动", "见第三章", "只总结有明确数据支撑的异常"),
    ]
    while len(table.rows) < len(rows):
        table.add_row()
    _clear_table(table)
    for r, row in enumerate(rows):
        for c, value in enumerate(row[: min(len(row), len(table.rows[r].cells))]):
            table.cell(r, c).text = value


def _fill_week_table(table, weeks: list[WeekStats], include_seo: bool) -> None:
    required_rows = 5 if include_seo else 5
    while len(table.rows) < required_rows:
        table.add_row()
    while len(table.columns) < len(weeks) + 1:
        table.add_column(table.columns[0].width)
    _clear_table(table)
    table.cell(0, 0).text = "来源"
    for idx, week in enumerate(weeks, start=1):
        table.cell(0, idx).text = f"{week.label}（{_short_range(week)}，{week.days}天）"
    rows = [
        ("AI 索引", lambda w: f"{w.ai_index:,} 次 / 日均 {w.daily_avg(w.ai_index):.1f} 次"),
        ("AI 搜索", lambda w: f"{w.ai_search:,} 次 / 日均 {w.daily_avg(w.ai_search):.1f} 次"),
        ("AI 训练", lambda w: f"{w.ai_training:,} 次 / 日均 {w.daily_avg(w.ai_training):.1f} 次"),
    ]
    if include_seo:
        rows.append(("SEO Bot", lambda w: f"{w.seo_bot:,} 次 / 日均 {w.daily_avg(w.seo_bot):.1f} 次"))
    for row_idx, (label, formatter) in enumerate(rows, start=1):
        table.cell(row_idx, 0).text = label
        for idx, week in enumerate(weeks, start=1):
            table.cell(row_idx, idx).text = formatter(week)
    _trim_table_rows(table, len(rows) + 1)


def _fill_shopify_compare_table(table, weeks: list[WeekStats], data: dict[str, Any]) -> None:
    while len(table.rows) < max(2, len(weeks) + 1):
        table.add_row()
    while len(table.columns) < 6:
        table.add_column(table.columns[0].width)
    _clear_table(table)
    headers = ["周期", "Bot Sessions", "占比", "Human Sessions", "占比", "日均总量"]
    for idx, value in enumerate(headers):
        table.cell(0, idx).text = value
    dashboard = data["dashboard_view"]
    total = dashboard["cards"]["total_requests"] or 1
    for row_idx, week in enumerate(weeks[:4], start=1):
        table.cell(row_idx, 0).text = f"{week.label} {_short_range(week)}"
        table.cell(row_idx, 1).text = f"{week.total:,}"
        table.cell(row_idx, 2).text = _share(week.ai_total, week.total)
        table.cell(row_idx, 3).text = f"{max(week.total - week.ai_total, 0):,}"
        table.cell(row_idx, 4).text = _share(max(week.total - week.ai_total, 0), week.total)
        table.cell(row_idx, 5).text = f"{week.daily_avg(week.total):.1f}"
    _trim_table_rows(table, len(weeks[:4]) + 1)


def _fill_non_shopify_compare_table(table, weeks: list[WeekStats], data: dict[str, Any]) -> None:
    while len(table.rows) < max(2, len(weeks) + 1):
        table.add_row()
    while len(table.columns) < 6:
        table.add_column(table.columns[0].width)
    _clear_table(table)
    headers = ["周期", "AI 总量", "AI 搜索", "AI 训练", "AI 索引", "说明"]
    for idx, value in enumerate(headers):
        table.cell(0, idx).text = value
    for row_idx, week in enumerate(weeks[:4], start=1):
        table.cell(row_idx, 0).text = f"{week.label} {_short_range(week)}"
        table.cell(row_idx, 1).text = f"{week.ai_total:,}"
        table.cell(row_idx, 2).text = f"{week.ai_search:,}"
        table.cell(row_idx, 3).text = f"{week.ai_training:,}"
        table.cell(row_idx, 4).text = f"{week.ai_index:,}"
        table.cell(row_idx, 5).text = "观察 AI 平均值与平台结构变化"
    _trim_table_rows(table, len(weeks[:4]) + 1)


def _fill_target_table(table, stage: dict[str, Any]) -> None:
    rows = [
        ("阶段目标", "参考数值", "意义"),
        ("当前水平", f"约 {stage.get('chatgpt_total', 0)} 次（ChatGPT-User）", "判断是否进入稳定推荐阶段"),
        ("短期目标", "连续 2 周稳定增长", "确认增长为可持续趋势"),
        ("中期目标", "AI 搜索访问继续扩张", "进入规模化 AI 推荐阶段"),
    ]
    while len(table.rows) < len(rows):
        table.add_row()
    _clear_table(table)
    for r, row in enumerate(rows):
        for c, value in enumerate(row):
            table.cell(r, c).text = value


def _fill_recommend_target_table(table, data: dict[str, Any]) -> None:
    rows = [
        ("建议方向", "参考指标", "意义"),
        ("整体流量追踪", f"{data['dashboard_view']['cards']['total_requests']:,} 次", "持续观察客户流量是否增长"),
        ("AI BOT构成", f"{data['summary']['ai_index']:,}/{data['summary']['ai_search']:,}/{data['summary']['ai_training']:,}", "跟踪 AI 索引/搜索/训练结构变化"),
        ("重点页面", "Top 5 页面定期复查", "判断 AI 对站点内容理解重点是否变化"),
    ]
    while len(table.rows) < len(rows):
        table.add_row()
    _clear_table(table)
    for r, row in enumerate(rows):
        for c, value in enumerate(row):
            table.cell(r, c).text = value


def _clear_table(table) -> None:
    for row in table.rows:
        for cell in row.cells:
            cell.text = ""


def _trim_table_rows(table, keep_rows: int) -> None:
    while len(table.rows) > keep_rows:
        table._tbl.remove(table.rows[-1]._tr)


def _detect_anomalies(data: dict[str, Any], shopify: bool) -> list[tuple[str, str]]:
    anomalies: list[tuple[str, str]] = []
    days = data["summary"]["days"]
    if not days:
        return anomalies
    totals = [row.get("ai_search", 0) + row.get("ai_training", 0) + row.get("ai_index", 0) + row.get("seo_bot", 0) for row in days]
    avg = sum(totals) / len(totals) if totals else 0
    peak_idx = max(range(len(totals)), key=lambda idx: totals[idx])
    if avg and totals[peak_idx] > avg * 1.8:
        anomalies.append((f"3.1  {days[peak_idx]['date']} 流量出现显著峰值", f"现象：单日总量 {totals[peak_idx]:,} 次，约为周期日均 {avg:.1f} 次的 {totals[peak_idx] / avg:.1f} 倍。"))
    if data["stage"].get("chatgpt_total", 0) > 0:
        anomalies.append(("3.2  ChatGPT-User 增长需要持续确认", f"现象：当前周期累计 ChatGPT-User {data['stage']['chatgpt_total']} 次，需继续观察是否保持稳定增长。"))
    rows = anomalies[:2]
    normalized: list[tuple[str, str]] = []
    for idx, (title, body) in enumerate(rows, start=1):
        fixed_title = re.sub(r"^3\.\d+\s*", f"3.{idx}  ", title, count=1)
        if fixed_title == title:
            fixed_title = f"3.{idx}  {title}"
        normalized.append((fixed_title, body))
    return normalized


def _fill_anomaly_section(doc, data: dict[str, Any], shopify: bool) -> None:
    paragraphs = doc.paragraphs
    llm = data.get("llm_sections", {}) or {}
    anomalies = _llm_anomalies(llm) or _detect_anomalies(data, shopify)
    if not anomalies:
        _set_paragraph_text(paragraphs[26], "三、待确认的异常")
        _set_paragraph_text(paragraphs[27], "本周期未发现需要重点确认的显著异常。")
        _set_paragraph_text(paragraphs[28], "")
        for idx in range(29, 38):
            _set_paragraph_text(paragraphs[idx], "" if idx not in (30, 35) else "—")
        return
    _set_paragraph_text(paragraphs[26], "三、待确认的异常")
    _set_paragraph_text(paragraphs[27], llm.get("anomaly_intro") or f"本章主动列出当前周期内需要进一步确认的异常，共 {len(anomalies)} 项。")
    _set_paragraph_text(paragraphs[28], "")
    for idx in range(29, 38):
        _set_paragraph_text(paragraphs[idx], "")
    title_slots = [29, 33]
    body_slots = [[30, 31, 32], [34, 35, 36]]
    for idx, (title, body) in enumerate(anomalies[:2]):
        _set_paragraph_text(paragraphs[title_slots[idx]], title)
        body_parts = _split_paragraphs(body, len(body_slots[idx]))
        for paragraph_idx, part in zip(body_slots[idx], body_parts):
            _set_paragraph_text(paragraphs[paragraph_idx], part)
    if len(anomalies) < 2:
        _set_paragraph_text(paragraphs[33], "")
        _set_paragraph_text(paragraphs[34], "—")
        _set_paragraph_text(paragraphs[35], "")
        _set_paragraph_text(paragraphs[36], "")
    _set_paragraph_text(paragraphs[37], "")


def _llm_anomalies(llm_sections: dict[str, Any]) -> list[tuple[str, str]]:
    if not isinstance(llm_sections, dict):
        return []
    rows: list[tuple[str, str]] = []
    seq = 1
    for idx in (1, 2):
        raw_title = str(llm_sections.get(f"anomaly_{idx}_title") or "").strip()
        body = str(llm_sections.get(f"anomaly_{idx}_body") or "").strip()
        if raw_title and body:
            title = re.sub(r"^3\.\d+\s*", f"3.{seq}  ", raw_title, count=1)
            if title == raw_title:
                title = f"3.{seq}  {raw_title}"
            rows.append((title, body))
            seq += 1
    return rows


def _split_paragraphs(text: str, parts: int) -> list[str]:
    chunks = [item.strip() for item in re.split(r"(?:\n+|(?<=。))", text) if item and item.strip()]
    if not chunks:
        return [""] * parts
    result: list[str] = []
    remaining = chunks[:]
    while remaining and len(result) < parts:
        slots_left = parts - len(result)
        take = max(1, len(remaining) - (slots_left - 1))
        result.append("".join(remaining[:take]).strip())
        remaining = remaining[take:]
    while len(result) < parts:
        result.append("")
    return result[:parts]


def _top_names(rows: list[dict[str, Any]] | None) -> str:
    if not rows:
        return "—"
    names = [row["name"] for row in rows[:4] if row.get("name")]
    return " / ".join(names) if names else "—"


def _overall_top_ai_platform(rankings: dict[str, list[dict[str, Any]]] | None) -> dict[str, Any]:
    rows = []
    for category in ("ai_search", "ai_training", "ai_index"):
        for item in (rankings or {}).get(category, []):
            rows.append(
                {
                    "name": str(item.get("name") or ""),
                    "requests": int(item.get("requests") or 0),
                    "category": category,
                }
            )
    rows = [row for row in rows if row["name"] and row["requests"] > 0]
    rows.sort(key=lambda row: (-row["requests"], row["name"]))
    return rows[0] if rows else {"name": "—", "requests": 0, "category": ""}


def _daily_avg(value: int, days: int) -> str:
    return f"{(value / days if days else 0):.1f}"


def _share(part: int, total: int) -> str:
    if not total:
        return "0%"
    return f"{(part * 100 / total):.1f}%"


def _pct_change(current: float, previous: float) -> str:
    if previous == 0:
        return "N/A"
    return f"{((current - previous) * 100 / previous):+.1f}%"


def _short_range(week: WeekStats) -> str:
    start = date.fromisoformat(week.start)
    end = date.fromisoformat(week.end)
    return f"{start.month}/{start.day}–{end.month}/{end.day}"


def _weeks_overview_text(weeks: list[WeekStats]) -> str:
    return "；".join(f"{week.label} = {week.start} 至 {week.end}（{week.days}天）" for week in weeks)


def _build_chart_images(data: dict[str, Any], output_dir: Path) -> dict[str, Path]:
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
    except ImportError as exc:
        raise DashboardQueryError(
            message="report generation dependency missing: matplotlib",
            error_type="report_dependency_missing",
            source="reporting",
            status_code=500,
            extra={"module": "matplotlib"},
        ) from exc

    output_dir.mkdir(parents=True, exist_ok=True)
    summary = data["summary"]
    days = summary["days"]
    weeks = data["week_stats"]
    rankings = data["rankings"]
    top_pages = data["top_pages"]
    dashboard_days = data.get("dashboard_days", []) or []
    is_shopify = bool(data.get("is_shopify"))
    colors = {
        "ai_index": "#4c6ef5",
        "ai_search": "#12b886",
        "ai_training": "#e0a458",
        "seo_bot": "#d16d5b",
    }

    paths: dict[str, Path] = {}

    def _annotate_bars(ax, bars, horizontal: bool = False, fmt: str = "{:,.0f}") -> None:
        for bar in bars:
            value = bar.get_width() if horizontal else bar.get_height()
            if abs(value) < 1e-9:
                continue
            if horizontal:
                ax.text(
                    bar.get_x() + bar.get_width() + max(abs(value) * 0.01, 0.3),
                    bar.get_y() + bar.get_height() / 2,
                    fmt.format(value),
                    va="center",
                    ha="left",
                    fontsize=8,
                )
            else:
                ax.text(
                    bar.get_x() + bar.get_width() / 2,
                    bar.get_y() + bar.get_height() + max(abs(value) * 0.01, 0.3),
                    fmt.format(value),
                    va="bottom",
                    ha="center",
                    fontsize=8,
                )

    # Figure 1
    fig, ax = plt.subplots(figsize=(10, 4.8), dpi=150)
    labels = [row["date"][5:] for row in days]
    total_values = [row.get("ai_index", 0) + row.get("ai_search", 0) + row.get("ai_training", 0) + (row.get("seo_bot", 0) if is_shopify else 0) for row in days]
    index_values = [row.get("ai_index", 0) for row in days]
    training_values = [row.get("ai_training", 0) for row in days]
    search_values = [row.get("ai_search", 0) for row in days]
    seo_values = [row.get("seo_bot", 0) for row in days]
    ax.plot(labels, total_values, marker="o", linewidth=2.4, color="#111827", label="Total" if is_shopify else "AI Total")
    ax.plot(labels, index_values, marker="o", linewidth=2.0, color=colors["ai_index"], label="AI Index")
    ax.plot(labels, training_values, marker="o", linewidth=2.0, color=colors["ai_training"], label="AI Training")
    ax.plot(labels, search_values, marker="o", linewidth=2.0, color=colors["ai_search"], label="AI Search")
    if is_shopify:
        ax.plot(labels, seo_values, marker="o", linewidth=2.0, color=colors["seo_bot"], label="SEO")
    ax.set_ylabel("Requests")
    ax.grid(alpha=0.18)
    ax.legend(frameon=False)
    plt.setp(ax.get_xticklabels(), rotation=90)
    visible_series = total_values + index_values + training_values + search_values
    if is_shopify:
        visible_series = visible_series + seo_values
    peak = max(visible_series + [1])
    floor = min(visible_series + [0])
    if is_shopify:
        headroom = max(peak * 0.16, 8)
        footroom = max((peak - floor) * 0.08, 4)
        ax.set_ylim(max(0, floor - footroom), peak + headroom)
    else:
        headroom = max(peak * 0.08, 4)
        ax.set_ylim(0, peak + headroom)
    ax.margins(x=0.05)
    fig.subplots_adjust(left=0.11, right=0.97, top=0.86, bottom=0.24)
    paths["word/media/image1.png"] = output_dir / "image1.png"
    fig.savefig(paths["word/media/image1.png"], bbox_inches="tight", pad_inches=0.24)
    plt.close(fig)

    # Figure 2
    fig_width = max(12, len(weeks) * 1.8)
    fig, axes = plt.subplots(1, 3, figsize=(fig_width, 4.4), dpi=150)
    week_labels = [week.label for week in weeks]
    palette = [colors["ai_index"], colors["seo_bot"], colors["ai_search"], colors["ai_training"], "#6b7280", "#8b5cf6"]
    bar_colors = [palette[i % len(palette)] for i in range(len(weeks))]
    bars0 = axes[0].bar(week_labels, [week.total for week in weeks], color=bar_colors)
    axes[0].set_title("Total")
    _annotate_bars(axes[0], bars0)
    bars1 = axes[1].bar(week_labels, [week.daily_avg(week.total) for week in weeks], color=bar_colors)
    axes[1].set_title("Daily Avg")
    _annotate_bars(axes[1], bars1, fmt="{:,.1f}")
    x = [0, 1] if is_shopify else [0, 1, 2]
    width = 0.8 / max(len(weeks), 1)
    for idx, week in enumerate(weeks):
        offset = -0.4 + width / 2 + idx * width
        values = [week.daily_avg(week.ai_total), week.daily_avg(week.seo_bot)] if is_shopify else [week.daily_avg(week.ai_search), week.daily_avg(week.ai_training), week.daily_avg(week.ai_index)]
        bars = axes[2].bar([i + offset for i in x], values, width=width, color=bar_colors[idx], label=week.label)
        _annotate_bars(axes[2], bars, fmt="{:,.1f}")
    axes[2].set_xticks(x)
    axes[2].set_xticklabels(["AI Total", "SEO Bot"] if is_shopify else ["AI Search", "AI Training", "AI Index"])
    axes[2].set_title("Category Avg")
    axes[2].legend(frameon=False)
    for ax in axes:
        ax.grid(axis="y", alpha=0.18)
        ax.margins(y=0.14)
        plt.setp(ax.get_xticklabels(), rotation=90)
    fig.subplots_adjust(left=0.06, right=0.98, top=0.88, bottom=0.18, wspace=0.28)
    paths["word/media/image2.png"] = output_dir / "image2.png"
    fig.savefig(paths["word/media/image2.png"], bbox_inches="tight", pad_inches=0.18)
    plt.close(fig)

    # Figure 3
    fig, ax = plt.subplots(figsize=(11, 8), dpi=150)
    page_rows = list(reversed(top_pages[:25]))
    names = [_trim_page_name(item["page"]) for item in page_rows]
    left = [0] * len(page_rows)
    page_series = [("ai_index", "AI Index"), ("ai_search", "AI Search"), ("ai_training", "AI Training")]
    if is_shopify:
        page_series.append(("seo_bot", "SEO Bot"))
    for key, label in page_series:
        vals = [item.get(key, 0) for item in page_rows]
        bars = ax.barh(names, vals, left=left, color=colors.get(key, "#999999"), label=label)
        left = [left[i] + vals[i] for i in range(len(vals))]
    for idx, total in enumerate(left):
        if total <= 0:
            continue
        ax.text(total + max(total * 0.01, 0.3), idx, f"{total:,.0f}", va="center", ha="left", fontsize=8)
    ax.grid(axis="x", alpha=0.18)
    ax.legend(frameon=False, loc="lower right")
    fig.subplots_adjust(left=0.22, right=0.98, top=0.96, bottom=0.06)
    paths["word/media/image3.png"] = output_dir / "image3.png"
    fig.savefig(paths["word/media/image3.png"], bbox_inches="tight", pad_inches=0.18)
    plt.close(fig)

    # Figure 4
    fig, ax = plt.subplots(figsize=(11, 4.8), dpi=150)
    compare_labels = [row["date"][5:] for row in dashboard_days] or labels
    compare_total = [row["total"] for row in dashboard_days] or total_values
    compare_ai = [row["ai_total"] for row in dashboard_days] or [row.get("ai_index", 0) + row.get("ai_search", 0) + row.get("ai_training", 0) for row in summary["days"]]
    base_series = compare_total if is_shopify else compare_ai
    compare_x = list(range(len(compare_labels)))
    bars = ax.bar(compare_x, base_series, color="#8aa29e", label="Total Requests" if is_shopify else "AI Requests")
    _annotate_bars(ax, bars)
    ax2 = ax.twinx()
    ax2.plot(compare_x, compare_ai, color="#d16d5b", marker="o", linewidth=2.2, label="AI Requests")
    for idx, value in enumerate(compare_ai):
        ax2.text(compare_x[idx], value + max(max(compare_ai + [1]) * 0.02, 0.5), f"{value:,.0f}", color="#d16d5b", ha="center", va="bottom", fontsize=8)
    ax.grid(axis="y", alpha=0.18)
    ax.set_ylabel("Total" if is_shopify else "AI")
    ax2.set_ylabel("AI")
    ax.set_xticks(compare_x)
    ax.set_xticklabels(compare_labels, rotation=90)
    lines, labels1 = ax.get_legend_handles_labels()
    lines2, labels2 = ax2.get_legend_handles_labels()
    ax.legend(lines + lines2, labels1 + labels2, frameon=False, loc="upper left")
    fig.subplots_adjust(left=0.08, right=0.92, top=0.90, bottom=0.20)
    paths["word/media/image4.png"] = output_dir / "image4.png"
    fig.savefig(paths["word/media/image4.png"], bbox_inches="tight", pad_inches=0.18)
    plt.close(fig)

    # Figure 5
    fig, ax = plt.subplots(figsize=(11, 4.8), dpi=150)
    compare_user = [row["user_total"] for row in dashboard_days] or [0 for _ in compare_labels]
    compare_bot = [max(row["total"] - row["user_total"], 0) for row in dashboard_days] or compare_total
    x = list(range(len(compare_labels)))
    width = 0.24 if not is_shopify else 0.38
    if is_shopify:
        bars_user = ax.bar([i - width / 2 for i in x], compare_user, width=width, color="#6c8c7f", label="User Requests")
        bars_bot = ax.bar([i + width / 2 for i in x], compare_bot, width=width, color="#caa96b", label="Bot Requests")
        _annotate_bars(ax, bars_user)
        _annotate_bars(ax, bars_bot)
        ax2 = ax.twinx()
        ax2.plot(x, compare_ai, color="#d16d5b", marker="o", linewidth=2.2, label="AI Requests")
        for idx, value in enumerate(compare_ai):
            ax2.text(idx, value + max(max(compare_ai + [1]) * 0.02, 0.5), f"{value:,.0f}", color="#d16d5b", ha="center", va="bottom", fontsize=8)
        ax.set_ylabel("User / Bot")
        ax2.set_ylabel("AI")
        lines, labels1 = ax.get_legend_handles_labels()
        lines2, labels2 = ax2.get_legend_handles_labels()
        ax.legend(lines + lines2, labels1 + labels2, frameon=False, loc="upper left")
    else:
        bars_index = ax.bar([i - width for i in x], [row.get("ai_index", 0) for row in dashboard_days], width=width, color=colors["ai_index"], label="AI Index")
        bars_search = ax.bar(x, [row.get("ai_search", 0) for row in dashboard_days], width=width, color=colors["ai_search"], label="AI Search")
        bars_training = ax.bar([i + width for i in x], [row.get("ai_training", 0) for row in dashboard_days], width=width, color=colors["ai_training"], label="AI Training")
        _annotate_bars(ax, bars_index)
        _annotate_bars(ax, bars_search)
        _annotate_bars(ax, bars_training)
        ax.set_ylabel("AI Requests")
        ax.legend(frameon=False, loc="upper left")
    ax.set_xticks(x)
    ax.set_xticklabels(compare_labels, rotation=90)
    ax.grid(axis="y", alpha=0.18)
    fig.subplots_adjust(left=0.08, right=0.92, top=0.90, bottom=0.20)
    paths["word/media/image5.png"] = output_dir / "image5.png"
    fig.savefig(paths["word/media/image5.png"], bbox_inches="tight", pad_inches=0.18)
    plt.close(fig)
    return paths


def _trim_page_name(value: str, limit: int = 42) -> str:
    return value if len(value) <= limit else value[: limit - 1] + "…"


def _replace_docx_images(docx_path: Path, replacements: dict[str, Path]) -> None:
    temp_path = docx_path.with_suffix(".tmp")
    with zipfile.ZipFile(docx_path, "r") as zin, zipfile.ZipFile(temp_path, "w") as zout:
        for info in zin.infolist():
            if info.filename in replacements:
                zout.writestr(info, replacements[info.filename].read_bytes())
            else:
                zout.writestr(info, zin.read(info.filename))
    temp_path.replace(docx_path)
